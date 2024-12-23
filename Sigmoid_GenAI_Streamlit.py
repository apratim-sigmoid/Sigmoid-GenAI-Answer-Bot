import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import re
import json
import time
from io import StringIO, BytesIO
import requests
from PIL import Image
from docx import Document
from docx.shared import Inches
import base64
# pip install python-docx

# Streamlit app configuration
st.set_page_config(
    page_title="Sigmoid GenAI Answer Bot",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

logo = Image.open("Images/sigmoid-logo.png")
st.image(logo, width=120)

# Custom CSS
st.markdown("""
    <style>
    .stAlert {
        padding: 1rem;
        margin-bottom: 1rem;
    }
    .st-emotion-cache-16idsys p {
        font-size: 1.1rem;
    }
    </style>
    """, unsafe_allow_html=True)

# Initialize session state
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'initialized' not in st.session_state:
    st.session_state.initialized = False
if 'current_data_source' not in st.session_state:
    st.session_state.current_data_source = None

def reset_app_state():
    """Reset the app state when data source changes"""
    st.session_state.initialized = False
    if 'df' in st.session_state:
        del st.session_state.df
        
def save_figure_to_image(fig):
    """Convert matplotlib figure to bytes for Word document."""
    buf = BytesIO()
    fig.savefig(buf, format='png', bbox_inches='tight')
    buf.seek(0)
    return buf

def create_word_document(chat_history):
    """Create a Word document from the analysis history."""
    doc = Document()
    doc.add_heading('Data Analysis History Report', 0)
    
    # Add generation timestamp
    doc.add_paragraph(f'Generated on: {time.strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph('-' * 50)
    
    # Add each analysis to the document
    for idx, chat in enumerate(reversed(chat_history), 1):
        # Add query section with data source
        doc.add_heading(f'Query: {chat["query"]}', level=1)
        doc.add_paragraph(f'Data Source: {chat.get("data_source", "Not specified")}')
        
        # Rest of the function remains the same
        if chat['approach']:
            doc.add_heading('Approach:', level=2)
            doc.add_paragraph(chat['approach'])
        
        if chat['answer']:
            doc.add_heading('Results:', level=2)
            doc.add_paragraph(chat['answer'])
        
        if chat['figure']:
            doc.add_heading('Visualization:', level=2)
            image_stream = save_figure_to_image(chat['figure'])
            doc.add_picture(image_stream, width=Inches(6))
        
        doc.add_paragraph('-' * 50)
    
    return doc

def download_word_doc():
    """Create and return a download link for the Word document."""
    if not st.session_state.chat_history:
        st.warning("No analysis history to export!")
        return
    
    # Create Word document
    doc = create_word_document(st.session_state.chat_history)
    
    # Save document to bytes
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    # Create download button
    st.download_button(
        label="📥 Download Analysis History",
        data=doc_bytes.getvalue(),
        file_name=f"analysis_history_{time.strftime('%Y%m%d_%H%M%S')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

def extract_code_segments(response_text):
    """Extract code segments from the API response using regex."""
    segments = {}
    
    # Extract approach section
    approach_match = re.search(r'<approach>(.*?)</approach>', response_text, re.DOTALL)
    if approach_match:
        segments['approach'] = approach_match.group(1).strip()
    
    # Extract content between <code> tags
    code_match = re.search(r'<code>(.*?)</code>', response_text, re.DOTALL)
    if code_match:
        segments['code'] = code_match.group(1).strip()
    
    # Extract content between <chart> tags
    chart_match = re.search(r'<chart>(.*?)</chart>', response_text, re.DOTALL)
    if chart_match:
        segments['chart'] = chart_match.group(1).strip()
    
    # Extract content between <answer> tags
    answer_match = re.search(r'<answer>(.*?)</answer>', response_text, re.DOTALL)
    if answer_match:
        segments['answer'] = answer_match.group(1).strip()
    
    return segments

def execute_analysis(df, response_text):
    """Execute the extracted code segments on the provided dataframe and store formatted answer."""
    results = {
        'approach': None,
        'answer': None,
        'figure': None,
        'code': None,
        'chart_code': None
    }
    
    try:
        # Extract code segments
        segments = extract_code_segments(response_text)
        
        if not segments:
            st.error("No code segments found in the response")
            return results
        
        # Store the approach and raw code
        if 'approach' in segments:
            results['approach'] = segments['approach']
        if 'code' in segments:
            results['code'] = segments['code']
        if 'chart' in segments:
            results['chart_code'] = segments['chart']
        
        # Create a single namespace for all executions
        namespace = {'df': df, 'pd': pd, 'plt': plt, 'sns': sns}
        
        # Execute analysis code and answer template
        if 'code' in segments and 'answer' in segments:
            # Properly dedent the code before execution
            code_lines = segments['code'].strip().split('\n')
            # Find minimum indentation
            min_indent = float('inf')
            for line in code_lines:
                if line.strip():  # Skip empty lines
                    indent = len(line) - len(line.lstrip())
                    min_indent = min(min_indent, indent)
            # Remove consistent indentation
            dedented_code = '\n'.join(line[min_indent:] if line.strip() else '' 
                                    for line in code_lines)
            
            # Combine code with answer template
            combined_code = f"""
{dedented_code}

# Format the answer template
answer_text = f'''{segments['answer']}'''
"""
            exec(combined_code, namespace)
            results['answer'] = namespace.get('answer_text')
        
        # Execute chart code if present
        if 'chart' in segments:
            # Properly dedent the chart code
            chart_lines = segments['chart'].strip().split('\n')
            # Find minimum indentation
            min_indent = float('inf')
            for line in chart_lines:
                if line.strip():  # Skip empty lines
                    indent = len(line) - len(line.lstrip())
                    min_indent = min(min_indent, indent)
            # Remove consistent indentation
            dedented_chart = '\n'.join(line[min_indent:] if line.strip() else '' 
                                     for line in chart_lines)
            
            plt.figure(figsize=(10, 6))
            exec(dedented_chart, namespace)
            fig = plt.gcf()
            results['figure'] = fig
            plt.close()
        
        return results
        
    except Exception as e:
        st.error(f"Error during execution: {str(e)}")
        return results
    

def get_prompt_file(data_source):
    """Return the appropriate prompt file based on the data source."""
    prompt_mapping = {
        'Outbound_Data.csv': 'Prompts/Prompt1.txt',
        'Inventory_Batch.csv': 'Prompts/Prompt2.txt',
        'Inbound_Data.csv': 'Prompts/Prompt3.txt'
    }
    return prompt_mapping.get(data_source)


def analyze_data_with_execution(df, question, api_key, data_source):
    # Get the appropriate prompt file based on data source
    prompt_file = get_prompt_file(data_source)
    
    if not prompt_file:
        st.error("Unable to determine prompt file for the selected data source!")
        return None

    # Read the prompt template from file
    try:
        with open(prompt_file, 'r') as file:
            data_description = file.read().strip()
    except FileNotFoundError:
        st.error(f"{prompt_file} file not found!")
        return None
    except Exception as e:
        st.error(f"Error reading {prompt_file}: {str(e)}")
        return None

    
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }
    
    payload = {
        "model": "gpt-4o",
        "messages": [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": f"""
                        
You are an AI assistant tasked with analyzing a dataset to provide code for calculating the final answer and generating relevant visualization.
I will provide you with the data in dataframe format, as well as a question to answer based on the data.

{data_description}

Here is the question I would like you to answer using this data:
<question>
{question}
</question>

To answer this, first think through your approach inside <approach> tags. Break down the steps you
will need to take and consider which columns of the data will be most relevant. Here is an example:
<approach>
To answer this question, I will need to:
1. Calculate the total number of orders and pallets across all rows
2. Determine the average distance and cost per order
3. Identify the most common PROD_TYPE and SHORT_POSTCODE
</approach>

Then, write the Python code needed to analyze the data and calculate the final answer inside <code> tags. Assume input dataframe as 'df'
Be sure to include any necessary data manipulation, aggregations, filtering, etc. Return only the Python code without any explanation or markdown formatting.
For decimal answers round them to 1 decimal place.

Generate Python code using matplotlib and/or seaborn to create an appropriate chart to visualize the relevant data and support your answer.
For example if user is asking for postcode with highest cost then a relevant chart can be a bar chart showing top 10 postcodes with highest total cost arranged in decreasing order.
Specify the chart code inside <chart> tags.
When working with dates:

Always convert dates to datetime using pd.to_datetime() with explicit format
For grouping by month, use dt.strftime('%Y-%m') instead of dt.to_period()
Sort date-based results chronologically before plotting

The visualization code should follow these guidelines:

Preferably use this color #d86a67 if possible, as it aligns with the color scheme of the dashboard

Start with these required imports:

import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns

Use standard chart setup:
# Set figure size and style
plt.figure(figsize=(8, 5))
# Set seaborn default style and color palette
sns.set_theme(style="whitegrid")  
sns.set_palette('pastel')

For time-based charts:


Use string dates on x-axis (converted using strftime)
Rotate labels: plt.xticks(rotation=45, ha='right')
Add gridlines: plt.grid(True, alpha=0.3)

For large numbers:
Format y-axis with K/M suffixes using:

Always include:

Clear title (plt.title())
Axis labels (plt.xlabel(), plt.ylabel())
plt.tight_layout() at the end


For specific chart types:

Time series: sns.lineplot() with marker='o'
Rankings: sns.barplot() with descending sort
Comparisons: sns.barplot() or sns.boxplot()
Distributions: sns.histplot() or sns.kdeplot()

Return only the Python code without any explanation or markdown formatting.

Finally, provide the answer to the question in natural language inside <answer> tags. Be sure to
include any key variables that you calculated in the code inside {{}}.
                    """
                    }
                ]
            }
        ],
        "max_tokens": 4096,
        "temperature": 0
    }

    try:
        response = requests.post("https://api.openai.com/v1/chat/completions", 
                              headers=headers, 
                              json=payload)
        
        if response.status_code != 200:
            st.error(f"Error: Received status code {response.status_code}")
            st.error(f"Response content: {response.text}")
            return None
        
        response_json = response.json()
        response_content = response_json['choices'][0]['message']['content']
        
        # Execute the code segments and get results
        results = execute_analysis(df, response_content)
        
        return results
        
    except Exception as e:
        st.error(f"Error during analysis: {e}")
        return None

    
def load_data_file(filename):
    """Load a CSV data file with automatic parsing of date columns."""
    try:
        # Load data without parsing dates first
        data = pd.read_csv(filename)
        
        # Identify columns with "date" in their name and parse them as dates
        date_columns = [col for col in data.columns if 'date' in col.lower()]
        
        # Reload data with date parsing for identified columns
        return pd.read_csv(filename, parse_dates=date_columns, dayfirst=True)
        
    except Exception as e:
        st.error(f"Error loading {filename}: {str(e)}")
        return None    

def display_analysis_results(results):
    """Display the analysis results in a structured format."""
    if results['approach']:
        st.subheader("Approach")
        st.write(results['approach'])
    
    if results['answer']:
        st.subheader("Analysis Results")
        st.write(results['answer'])
    
    if results['figure']:
        st.subheader("Visualization")
        st.pyplot(results['figure'])
    
    # Display code segments in expandable sections
    if results['code'] or results['chart_code']:
        st.subheader("Code Segments")
        
        if results['code']:
            with st.expander("Show Analysis Code"):
                st.code(results['code'], language='python')
        
        if results['chart_code']:
            with st.expander("Show Visualization Code"):
                st.code(results['chart_code'], language='python')
                
                
def get_sample_queries(data_source):
    """Return appropriate sample queries based on the selected data source."""
    queries = {
        'Outbound_Data.csv': [
            "Plot the bar chart for short postcode 'LU' in which x axis is number of pallets and y axis is total orders. Stack them on the basis of prod type.",
            "Which postcode results in the highest total cost?",
            "What is the monthly trend in total cost?",
            "What is the average cost per pallet for each PROD TYPE and how does it vary across the following SHORT_POSTCODE regions: CV, NG, NN, RG?",
            "Identify the distribution of cost per pallet, is it normally distributed?",
            "Generate a radar chart of average pallets per order for the top 15 postcodes with maximum average cost per order.",
            "Generate the boxplot distribution for pallets of the top 8 customers by total orders.",
            "For ambient product type, which are the top 5 customers with total orders > 10 and highest standard deviation in cost per pallet?",
            "What is the trend in cost over time and plot forecasted cost using 3-month exponential smoothing?",
            "Perform a hypothesis test to analyze if average cost per order differs significantly with product type.",
            "Create a regression line for cost per order and distance along with R squared.",
            "What is the distribution of cost in percentiles?",
            "How does the cost per order vary with distance within each PROD TYPE?",
            "Find the top 5 customers by total pallets shipped and compare their average cost per pallet and distance traveled.",
            "Identify the SHORT_POSTCODE areas with the highest total shipping costs and also mention their cost per pallet.",
            "Which customer has the highest total shipping cost over time, and how does its cost trend vary by month?",
            "What is the order frequency per week for the last 2 months?",
            "What is the total cost for ambient product type in January 2024?",
            "How has the cost per pallet evolved over the last 3 months?",
            "What is the average cost per pallet for each product type?"
        ],
        'Inventory_Batch.csv': [
            "What is the total inventory value by product category?",
            "Which products have inventory levels below their safety stock?",
            "What is the monthly trend in inventory turnover rate?",
            "Show the age distribution of current inventory batches",
            "Which are the top 10 products by storage cost?",
            "What is the average shelf life remaining for each product category?",
            "Identify products with excess inventory (more than 120% of max stock level)",
            "What is the weekly trend in inventory receipts vs. withdrawals?",
            "Generate a heat map of inventory levels across different storage locations",
            "Which products have the highest holding costs in the last quarter?",
            "Show the distribution of batch sizes by product category",
            "What is the correlation between product value and storage duration?",
            "Identify seasonal patterns in inventory levels for the top 5 products",
            "Calculate and visualize the inventory accuracy rates by location",
            "What is the average time between receipt and first withdrawal for each product?",
            "Show the distribution of inventory value across different temperature zones",
            "Which products have the highest stock rotation frequency?",
            "Generate a Pareto chart of inventory value by product category",
            "What is the trend in average days of inventory on hand?"
        ],
        'Inbound_Data.csv': [
            "What is the utilization in each tradelane for top 15 tradelane by pallets?",
            "What is the total cost in each route from Nov 2023 to Jan 2024? Consider top 10 routes with highest total pallets.",
            "What is the monthly trend of above metrics?",
            "What is the cost breakdown by Company?",
            "What is the proportion of FTL/LTL by route?",
            "What is the Pallet per Order?",
            "What is the cost per pallet?",
            "What is the cost per order?",
            "What is the average lead time by tradelane/tradeline/route?",
            "Which routes/delivery supplier/delivery groups charge higher fuel costs?",
            "Which routes/delivery supplier/delivery groups have higher % of late delivery?",
            "Which routes/delivery supplier/delivery groups have higher % of late collection?",
            "What is the average delay in delivery on a particular route by delivery supplier?",
            "What is the average delay in collection on a particular route by delivery supplier?"
        ]      
    }
    return queries.get(data_source, [])
                

def main():
    st.title("GenAI Answer Bot")
    
    # Define available data files
    data_files = {
        'Outbound_Data.csv': 'Data/Outbound_Data.csv',
        'Inventory_Batch.csv': 'Data/Inventory_Batch.csv',
        'Inbound_Data.csv': 'Data/Inbound_Data.csv'
    }
    
    # Sidebar configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # API Key input
        st.subheader("1. API Key")
        api_key = st.text_input("Enter OpenAI API Key:", type="password")
        
        # Data source selection
        st.subheader("2. Data Source")
        data_source = st.radio(
            "Choose Data Source:",
            # list(data_files.keys()) + ["Upload Custom File"],
            list(data_files.keys()),
            disabled=False,
            index=0
        )
        
        # Reset state if data source changes
        if st.session_state.current_data_source != data_source:
            st.session_state.current_data_source = data_source
            reset_app_state()
        
        df = None
        if data_source in data_files:
            df = load_data_file(data_files[data_source])
            if df is not None:
                st.success(f"{data_source} loaded successfully!")
                st.session_state.df = df
        else:
            uploaded_file = st.file_uploader("Upload CSV file", type=['csv'])
            if uploaded_file:
                try:
                    temp_df = pd.read_csv(uploaded_file)

                    # Identify columns with "date" in their name and parse those as dates
                    date_columns = [col for col in temp_df.columns if 'date' in col.lower()]
                    
                    # Reload the data with date parsing for identified columns
                    df = pd.read_csv(uploaded_file, parse_dates=date_columns, dayfirst=True)

                    st.success("Custom file loaded successfully!")
                    st.session_state.df = df
                except Exception as e:
                    st.error(f"Error loading custom file: {str(e)}")
    
    # Main content area
    if not api_key:
        st.info("Please enter your OpenAI API key in the sidebar to get started.")
        return
    
    if 'df' not in st.session_state:
        if data_source in data_files:
            st.error(f"Data file not found. Please check if '{data_source}' exists.")
        else:
            st.info("Please upload your CSV file in the sidebar.")
        return
    
    # Display sample data
    with st.expander("📊 View Sample Data"):
        display_df = st.session_state.df.copy()
        
        # Identify and format all datetime columns
        date_columns = display_df.select_dtypes(include=['datetime64']).columns
        for date_col in date_columns:
            display_df[date_col] = display_df[date_col].dt.strftime('%d-%m-%Y')
        
        display_df = display_df.set_index(display_df.columns[0])    
        st.dataframe(display_df.head(), use_container_width=True)
    
    
    # Query interface
    st.subheader("💬 Ask Questions About Your Data")
    
    # Get sample queries based on selected data source
    sample_queries = get_sample_queries(st.session_state.current_data_source)
    
    selected_query = st.selectbox(
        "Select a sample query or write your own below:",
        [""] + sample_queries,
        key="query_select"
    )
    
    query = st.text_area(
        "Enter your query:",
        value=selected_query,
        height=100,
        key="query_input"
    )
    
    col1, col2 = st.columns([1, 5])
    with col1:
        submit_button = st.button("🔍 Analyze")
    
    if submit_button and query:
        # Move time tracking and spinner to encompass both analysis and display
        with st.spinner("Analyzing data and generating visualizations..."):
            start_time = time.time()
            
            # Perform analysis
            results = analyze_data_with_execution(
                st.session_state.df, 
                query, 
                api_key, 
                st.session_state.current_data_source
            )
            
            if results:
                # Display results inside the spinner context
                display_analysis_results(results)
                
                # Store in chat history
                chat_entry = {
                    "data_source": st.session_state.current_data_source,
                    "query": query,
                    "approach": results['approach'],
                    "answer": results['answer'],
                    "figure": results['figure'],
                    "code": results['code'],
                    "chart_code": results['chart_code'],
                }
                
                # Check if this exact query isn't already the last entry
                if not st.session_state.chat_history or st.session_state.chat_history[-1]["query"] != query:
                    st.session_state.chat_history.append(chat_entry)
                
                end_time = time.time()
                time_taken = end_time - start_time
                
        # Show completion message after spinner
        st.info(f"Analysis completed in {time_taken:.1f} seconds")
    
    
    # Display analysis history with download and delete options
    if st.session_state.chat_history:
        col1, col2 = st.columns([6, 2])
        with col1:
            st.subheader("📜 Analysis History")
        with col2:
            download_word_doc()
            
        # Iterate through history in reverse order with index tracking
        for idx, chat in enumerate(reversed(st.session_state.chat_history)):
            # Calculate the actual index in the original list
            original_idx = len(st.session_state.chat_history) - idx - 1
            
            # Create two columns for each analysis entry
            hist_col1, hist_col2 = st.columns([20, 1])
            
            with hist_col1:
                with st.expander(
                    f"Query {len(st.session_state.chat_history) - idx}: {chat['query'][:50]}...",
                    expanded=False
                ):
                    st.markdown("**🔍 Query:**")
                    st.write(chat['query'])
                    
                    st.markdown(f"**📊 Data Source:** {chat.get('data_source', 'Not specified')}")
                    
                    if chat['approach']:
                        st.markdown("**🎯 Approach:**")
                        st.write(chat['approach'])
                    
                    if chat['answer']:
                        st.markdown("**💡 Results:**")
                        st.write(chat['answer'])
                    
                    if chat['figure']:
                        st.pyplot(chat['figure'])
            
            with hist_col2:
                # Add delete button for each entry
                if st.button("🗑️", key=f"delete_{original_idx}"):
                    st.session_state.chat_history.pop(original_idx)
                    st.rerun()  # Rerun the app to refresh the display
                    

if __name__ == "__main__":
    main()